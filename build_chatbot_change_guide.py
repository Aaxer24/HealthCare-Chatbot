from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = Path("Healthcare_Chatbot_Change_Guide.docx")

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
LIGHT_FILL = "F2F4F7"
TEXT = RGBColor(0x18, 0x24, 0x33)
MUTED = RGBColor(0x5B, 0x67, 0x76)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_inches):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))


def apply_font(run, name="Calibri", size=11, bold=False, color=TEXT, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Callout" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.5)
        style.font.color.rgb = TEXT
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.left_indent = Inches(0.15)
        style.paragraph_format.right_indent = Inches(0.15)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Healthcare Chatbot Change Guide")
    apply_font(run, size=24, bold=True, color=TEXT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(
        "A detailed explanation of every major change made to the project, why it was made, "
        "how the implementation evolved, and what each decision teaches about building a "
        "production-grade medical RAG chatbot."
    )
    apply_font(run, size=11, color=MUTED)

    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    widths = [1.7, 4.8]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    labels = [("Project", "HealthCare Chatbot"), ("Scope", "Architecture, retrieval, prompting, UI, routing, safety, and deployment hardening")]
    for row, pair in zip(table.rows, labels):
        set_cell_shading(row.cells[0], LIGHT_FILL)
        left = row.cells[0].paragraphs[0].add_run(pair[0])
        apply_font(left, size=10.5, bold=True)
        right = row.cells[1].paragraphs[0].add_run(pair[1])
        apply_font(right, size=10.5)

    doc.add_paragraph()


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_width(cell, 6.3)
    set_cell_shading(cell, "F7FBFF")
    p = cell.paragraphs[0]
    title_run = p.add_run(title + ": ")
    apply_font(title_run, size=10.5, bold=True, color=DARK_BLUE)
    body_run = p.add_run(body)
    apply_font(body_run, size=10.5)
    doc.add_paragraph()


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent = Inches(0.25 + (0.2 * level))
    p.paragraph_format.first_line_indent = Inches(-0.18)
    run = p.add_run(u"\u2022 " + text)
    apply_font(run)


def add_numbered_item(doc, title, body):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(6)
    title_run = p.add_run(title + " ")
    apply_font(title_run, bold=True)
    body_run = p.add_run(body)
    apply_font(body_run)


def add_section(doc, heading, paragraphs):
    doc.add_paragraph(heading, style="Heading 1")
    for para in paragraphs:
        p = doc.add_paragraph(style="Normal")
        run = p.add_run(para)
        apply_font(run)


def add_change_table(doc):
    doc.add_paragraph("High-level change map", style="Heading 1")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    headers = ["Area", "What changed", "Why it mattered"]
    widths = [1.5, 2.5, 2.5]
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, LIGHT_FILL)
        run = cell.paragraphs[0].add_run(headers[idx])
        apply_font(run, size=10.5, bold=True)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    rows = [
        ("Project structure", "Split one file into modular src package", "Made the code maintainable, explainable, and easier to extend"),
        ("Medical prompt", "Rewrote answer behavior and safety instructions", "Improved answer tone, structure, and trustworthiness"),
        ("General chat routing", "Added routing for greetings and chatbot-only questions", "Stopped simple chat from going through medical retrieval"),
        ("Out-of-scope handling", "Added redirect for random non-medical questions", "Kept the chatbot aligned to medical use instead of acting like a generic assistant"),
        ("Embeddings", "Moved from MiniLM to BGE small", "Improved retrieval quality for semantically similar medical text"),
        ("Reranking", "Added cross-encoder reranking", "Improved relevance of final retrieved passages"),
        ("Source previews", "Added expandable citation snippets", "Improved explainability and user trust"),
        ("UI and styling", "Redesigned the Streamlit interface", "Made the app feel more polished and production-like"),
        ("Typing effect", "Added streamed word-by-word rendering", "Made responses feel more conversational"),
        ("Config and docs", "Added env template, README, metadata checks", "Improved reproducibility and deployment hygiene"),
    ]
    for row_data in rows:
        row = table.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            run = cell.paragraphs[0].add_run(text)
            apply_font(run, size=10.3)

    doc.add_paragraph()


def main():
    doc = Document()
    style_document(doc)
    add_title(doc)
    add_callout(
        doc,
        "Purpose of this guide",
        "This document explains the chatbot's evolution from a basic Streamlit RAG prototype "
        "into a more production-oriented medical assistant. It is written as a learning guide, "
        "so each change includes both implementation details and the reasoning behind it.",
    )

    add_section(
        doc,
        "1. What the original chatbot looked like",
        [
            "At the beginning, the project was centered around a single main Streamlit file and a simple FAISS-based retrieval flow. The chatbot could load a local vector database, send the top chunks to a Groq model, and answer a user's question with basic citations.",
            "That early version was useful as a proof of concept, but it had the typical limitations of a first RAG prototype. Most responsibilities lived in one file, the prompting was rigid, the retrieval settings were basic, source presentation was rough, and user experience details were not yet tuned for a production-facing application.",
            "The initial code also treated almost every message as a retrieval question. That meant even simple messages like greetings were pushed through the medical RAG pipeline, which made the chatbot feel robotic and occasionally strange.",
        ],
    )

    add_change_table(doc)

    add_section(
        doc,
        "2. Why the refactor started with architecture",
        [
            "One of the earliest priorities was making the project easier to reason about. A single large file is manageable at the prototype stage, but once routing logic, prompt logic, retrieval upgrades, UI changes, and deployment details begin to accumulate, the code becomes harder to debug and harder to explain.",
            "To solve that, the app was reorganized into a small src package. Configuration moved into config.py, prompts moved into prompts.py, retrieval and reranking moved into rag.py, message routing moved into router.py, and Streamlit presentation logic moved into ui.py. The main medibot.py file became a thin entrypoint that coordinates those modules.",
            "This change matters because architecture is not just about neatness. A modular design reduces accidental coupling. It also makes future changes safer. For example, changing retrieval settings no longer requires editing UI code, and changing the prompt system no longer risks breaking layout behavior.",
        ],
    )

    add_callout(
        doc,
        "Design principle",
        "The refactor was driven by separation of concerns. Each module now owns one layer of behavior: configuration, prompt policy, retrieval, routing, or presentation.",
    )

    add_section(
        doc,
        "3. How the answer behavior became more human and more professional",
        [
            "The original prompt produced answers in a fixed template with sections like 'Short answer', 'What the documents say', and 'What to do next'. While this structure was predictable, it often sounded stiff. It also made simple answers feel inflated and made the assistant sound less like a real healthcare information guide and more like a mechanical report generator.",
            "The prompt system was rewritten so the model could answer naturally while still staying grounded. Instead of enforcing the same output skeleton every time, the prompt now tells the model to start with a direct answer, then expand only when the question calls for it. This keeps short questions concise and allows explanation-focused questions to become fuller and more educational.",
            "A second layer was added to adapt answer depth. Before medical generation, the app asks the model to classify whether the user likely wants a concise answer, an explanation, or practical guidance. That small routing step produces a style instruction that is appended to the medical question. The result is not hardcoded wording, but dynamic control over the level of detail.",
        ],
    )

    add_numbered_item(
        doc,
        "Why this matters.",
        "In a healthcare assistant, tone is part of trust. If the answer sounds too robotic, users often feel the system is brittle. If it sounds too casual, users may doubt its seriousness. The new prompt aims for a middle ground: calm, clear, professional, and human.",
    )

    add_section(
        doc,
        "4. Safety and scope controls",
        [
            "A medical chatbot needs better boundaries than a generic chatbot. Several changes were made to define what the system should and should not do. The safety prompt now explicitly tells the model not to diagnose, not to prescribe treatment, not to replace a clinician, and to direct urgent symptom patterns toward emergency care.",
            "The other important boundary is scope. During development, a general-chat path was added so greetings and simple conversational messages would not run through retrieval. That was useful, but it became too broad at one point and allowed unrelated questions such as pop-culture or sports questions to be answered as if the bot were a general assistant.",
            "That behavior was later tightened. The routing system now uses three buckets: simple chatbot conversation, real medical questions, and out-of-scope non-medical questions. If the user asks something unrelated to healthcare, the assistant now redirects them back toward medical use rather than inventing a general-domain answer.",
        ],
    )

    add_section(
        doc,
        "5. Why the retrieval pipeline was upgraded",
        [
            "The first retrieval setup used all-MiniLM-L6-v2 embeddings and straightforward FAISS lookup. That combination is excellent for a fast proof of concept, but it has limits when the corpus becomes larger or when medical language becomes more nuanced. Medical text contains synonyms, layered context, and subtle distinctions in symptoms, conditions, and recommendations.",
            "To improve retrieval quality, the embedding model was upgraded to BAAI/bge-small-en-v1.5. This model generally produces stronger semantic representations than the previous MiniLM embedding model. The benefit is that related medical passages are more likely to be retrieved even when the user's wording differs from the source text.",
            "Retrieval was also changed to use MMR, which stands for maximal marginal relevance. Instead of returning only the most similar chunks, MMR tries to balance relevance with diversity. That helps reduce repeated chunks and can produce a more useful context set for generation.",
        ],
    )

    add_callout(
        doc,
        "Tradeoff",
        "Better embeddings and MMR usually improve answer grounding, but they can increase latency and require a one-time rebuild of the vectorstore because old embeddings are no longer compatible with the new model.",
    )

    add_section(
        doc,
        "6. Why reranking was added after retrieval",
        [
            "Approximate vector search is good at finding candidate passages, but it is not always perfect at ranking the very best chunk first. To improve the last-mile quality of retrieval, a reranking stage was introduced using a cross-encoder model: cross-encoder/ms-marco-MiniLM-L-6-v2.",
            "The logic now works in two stages. First, FAISS plus MMR retrieves a broader candidate set. Second, the cross-encoder scores each candidate passage in the context of the exact user query. Because the cross-encoder reads the pair together, it is often better at deciding which passages are actually the most relevant to answer generation.",
            "This is one of the most production-oriented upgrades in the project. It pushes the system closer to a retrieval stack used in more mature RAG systems. The downside is speed. Reranking adds another model pass, which is one reason the app feels slower now than it did earlier.",
        ],
    )

    add_section(
        doc,
        "7. What changed in source citations and why",
        [
            "The earlier chatbot could show source file names, but the presentation was basic and not especially useful for users trying to understand why the answer should be trusted. To improve explainability, citations were upgraded into source previews.",
            "The retrieval layer now formats cleaner source labels, carries page information, and extracts a compact snippet from the retrieved chunk. In the UI, these snippets appear inside an expandable reference section so users can inspect the evidence without cluttering the main answer area.",
            "This change is important because medical trust depends on traceability. A strong answer becomes much more persuasive when the user can see where it came from and what part of the document supported it.",
        ],
    )

    add_section(
        doc,
        "8. Changes to general conversation behavior",
        [
            "Originally, greetings were hardcoded. That solved the immediate problem of the bot answering 'hi' with a medical-style response, but it was too brittle. Natural language always produces variants, and hardcoding quickly becomes messy.",
            "The project then moved to an LLM-based routing approach so the chatbot could decide whether a message was medical, conversational, or outside scope. This made the behavior more flexible and reduced the amount of literal phrase matching in code.",
            "However, the first version of that routing was too permissive. It allowed messages like 'what is wwe' to be answered, which went beyond the intended scope. The router was refined so only simple conversational or chatbot-identity messages stay in the general-chat path. Everything else either goes to medical retrieval or receives a polite scope reminder.",
        ],
    )

    add_section(
        doc,
        "9. User interface improvements and why they matter",
        [
            "The original Streamlit interface was functional, but it looked like a prototype. Styling was basic, spacing was uneven, citations were visually awkward, and some elements such as feedback buttons briefly appeared in intrusive ways.",
            "The interface was redesigned with a clearer hero area, sidebar status cards, cleaner chat panels, a more polished chat input, and a softer visual hierarchy. The result is still a Streamlit app, but it is now closer to a purposeful product interface instead of a default demo layout.",
            "The chat rendering was also improved with a word-by-word streaming effect. At first, the streamed text sometimes looked oversized because partial Markdown headings were being rendered while the answer was still incomplete. That was fixed by temporarily rendering the live stream as plain escaped text and only applying full Markdown formatting once the answer finished.",
            "Another UI cleanup involved removing the Good and Bad feedback buttons from the main chat surface. They were initially added as a production-oriented evaluation feature, but in practice they made the interface noisier and appeared in awkward positions, so they were removed to keep the experience focused.",
        ],
    )

    add_section(
        doc,
        "10. Model and configuration changes",
        [
            "The Groq model configuration also needed correction during development. One early default pointed to a model identifier that the current account could not access, which caused a model-not-found error. The default was updated to a Groq-supported model that works more reliably for this setup.",
            "Configuration handling was moved into environment-backed settings so the project can be run more predictably across environments. The app now reads values such as model name, retrieval depth, rerank depth, and feature flags from environment variables or Streamlit secrets.",
            "An .env.example file was added so the expected configuration is documented. This is a small but very real production improvement. A project becomes much easier to run, share, and deploy when the required configuration surface is explicit instead of hidden inside code.",
        ],
    )

    add_section(
        doc,
        "11. Vectorstore hygiene and reproducibility",
        [
            "When the embedding model changed, the previously built FAISS index became conceptually stale because it had been created with a different embedding space. To prevent subtle bugs, vectorstore metadata was introduced.",
            "The indexing script now saves a metadata file alongside the FAISS store, recording the embedding model and chunking details used during index construction. On app startup, the retrieval layer checks this metadata. If the vectorstore is missing or was built with a different embedding model, the app raises a clear rebuild instruction instead of failing silently or returning low-quality retrieval results.",
            "This kind of metadata guard is a production-minded detail. It protects the system from configuration drift and helps future maintainers understand when they need to rebuild the index.",
        ],
    )

    add_section(
        doc,
        "12. Supporting scripts and documentation improvements",
        [
            "The supporting scripts were upgraded to match the main application architecture. The vector-building script now uses the shared configuration and improved embedding settings. The CLI testing script also reuses the production retrieval logic so command-line tests resemble real app behavior.",
            "The repository gained a cleaner README, a safer .gitignore, and a more accurate setup story. These changes do not change answer quality directly, but they dramatically improve project usability. A production-grade application is not only one that runs well. It is also one that another person can understand, install, and operate without confusion.",
        ],
    )

    doc.add_paragraph("13. End-to-end request flow in the current system", style="Heading 1")
    flow_steps = [
        "The user sends a message in the Streamlit interface.",
        "The app renders the user message immediately and decides how to route the request.",
        "A lightweight classifier labels the message as GENERAL_CHAT, MEDICAL_QUESTION, or OUT_OF_SCOPE.",
        "If the message is GENERAL_CHAT, the assistant replies conversationally without retrieval.",
        "If the message is OUT_OF_SCOPE, the assistant politely redirects the user back to medical use.",
        "If the message is MEDICAL_QUESTION, a second style router decides whether the user wants a concise answer, an explanation, or practical guidance.",
        "The medical query then goes through FAISS retrieval with MMR, followed by optional reranking with the cross-encoder.",
        "The selected document chunks are passed to the Groq model with the medical system prompt and style instruction.",
        "The answer is streamed into the UI and any source previews are shown in the reference expander.",
    ]
    for step in flow_steps:
        add_bullet(doc, step)
    doc.add_paragraph()

    add_section(
        doc,
        "14. Why the app became slower and what that teaches",
        [
            "A useful lesson from this project is that quality improvements often increase latency. The original prototype took a short path from question to retrieval to answer. The upgraded system now performs message classification, answer-style routing, better retrieval, reranking, and then the final generation step, followed by streamed display.",
            "That extra latency is not an accident. It is the cost of stronger behavior. The system is more accurate, better scoped, more explainable, and more professional. Still, it is important to understand the tradeoff because production engineering is often about balancing quality, trust, and speed rather than maximizing only one of them.",
            "This is why later conversation in the project began to focus on which features should remain dynamic and which should be simplified. For example, general-chat handling only needs to cover greetings and chatbot-related small talk, so broad non-medical intelligence is not worth extra complexity here.",
        ],
    )

    add_section(
        doc,
        "15. What the project teaches overall",
        [
            "The evolution of this chatbot is a good example of how real applications mature. At first, the goal is proving that the idea works. After that, the work shifts toward reliability, boundaries, maintainability, retrieval quality, user trust, and UX polish.",
            "The most important lesson is that every change should connect to a concrete problem. Prompts were changed because answers sounded robotic. Routing was added because greetings should not hit retrieval. Scope controls were tightened because the assistant should not drift into unrelated domains. Better embeddings and reranking were added because retrieval quality matters more than just producing an answer quickly. UI changes mattered because production apps need credibility as well as correctness.",
            "In that sense, the chatbot is no longer just a small demo. It now reflects a more mature engineering mindset: one that treats architecture, safety, explainability, and user experience as part of the product rather than optional extras.",
        ],
    )

    doc.add_paragraph("Appendix A. File-by-file summary", style="Heading 1")
    appendix_rows = [
        ("medibot.py", "Converted into a thin Streamlit entrypoint that wires together configuration, routing, retrieval, and UI."),
        ("src/config.py", "Centralized runtime configuration, paths, model defaults, and feature flags."),
        ("src/prompts.py", "Holds the medical system prompt, follow-up condense prompt, general-chat classifier prompt, out-of-scope prompt, and answer-style prompt."),
        ("src/router.py", "Uses the LLM to classify messages, answer simple chatbot conversation, redirect out-of-scope questions, and decide answer depth."),
        ("src/rag.py", "Owns embeddings, FAISS loading, reranking, source formatting, snippet extraction, and medical answer generation."),
        ("src/ui.py", "Owns styling, header rendering, sidebar rendering, empty state, and streaming UI helpers."),
        ("create_memory_for_llm.py", "Builds the vectorstore with the newer embedding model and writes metadata for rebuild safety."),
        ("create_memory_with_llm.py", "Provides a command-line smoke test aligned to the production retrieval stack."),
        ("README.md", "Documents setup, architecture, features, and operational expectations."),
    ]
    for file_name, desc in appendix_rows:
        add_numbered_item(doc, file_name, desc)

    doc.add_paragraph("Appendix B. Suggested next learning steps", style="Heading 1")
    next_steps = [
        "Compare the original single-file version and the current modular version to see how responsibilities were separated.",
        "Read the prompts in src/prompts.py and map each one to the behavior it controls in the app.",
        "Trace a medical question through medibot.py, router.py, and rag.py to understand the live request flow.",
        "Experiment with retrieval_k and rerank_k in the environment configuration and observe the effect on speed and answer quality.",
        "Rebuild the vectorstore with additional medical PDFs and inspect how source previews change as the knowledge base grows.",
    ]
    for step in next_steps:
        add_bullet(doc, step)

    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
